import re
import tempfile
import zipfile
from collections import OrderedDict
from copy import deepcopy

import xlrd
from docx import Document
from lxml import etree

from ConfigurableSpiders.supports.tools import ArrayHelper, StringHelper, FileHelper, LxmlHelper


class StructuringParser:
    __tmp_dir_path = FileHelper.get_project_root_path()  # 临时文件夹，用于储存转换的临时文件

    @staticmethod
    def parse(response):
        content_type = None
        # 重定向后的response没有Content-Type
        if b'Content-Type' in response.headers:
            content_type = response.headers[b'Content-Type']
        postfix = FileHelper.get_postfix(response.request.url)
        if content_type is not None and b'text/html' in content_type:
            # html
            return StructuringParser.parse_html(response.body.decode(response.encoding))
        elif postfix == "docx":
            # docx文件
            file = tempfile.TemporaryFile()
            file.write(response.body)
            return StructuringParser.parse_docx(file)
        elif postfix == "doc":
            # doc文件
            doc_file_path = FileHelper.create_temp_file_by_bytes(response.body, "doc")
            try:
                file = FileHelper.convert_file_to_docx(doc_file_path)
            except Exception as e:
                print(response.request.url+"转换为docx文件失败")
                print(e)
                return []
        elif postfix == "xls" or postfix == 'xlsx':
            # xls文件
            file = FileHelper.create_temp_file_by_bytes(response.body, "xls")
        elif postfix == 'pdf':
            # pdf文件
            pdf_file_path = FileHelper.create_temp_file_by_bytes(response.body, "pdf")
            file = FileHelper.convert_file_to_docx(pdf_file_path)
        elif postfix == 'zip':
            # zip文件
            zip_file_path = FileHelper.create_temp_file_by_bytes(response.body, "zip")
            file = FileHelper.unpack_zip(zip_file_path)
        else:
            print('[error] 未处理的文件类型:contentType:' + str(content_type) + ",后缀:" + postfix)
            return []
        if isinstance(file, list):
            result = []
            for x in file:
                result.append(StructuringParser.dispatch_file_handler(x))
            return result
        else:
            return StructuringParser.dispatch_file_handler(file)

    @staticmethod
    def dispatch_file_handler(file_path):
        postfix = FileHelper.get_postfix(file_path)
        if postfix == "docx":
            # docx文件
            return StructuringParser.parse_docx(file_path)
        elif postfix == "doc":
            # doc文件
            return StructuringParser.parse_docx(file_path)
        elif postfix == "xls" or postfix == 'xlsx':
            # xls文件
            return StructuringParser.parse_xls(file_path)
        elif postfix == 'pdf':
            # pdf文件
            return StructuringParser.parse_docx(file_path)
        print("未处理的文件类型:" + postfix)

    @staticmethod
    def parse_html(html_str):
        doc = etree.HTML(html_str)
        body = doc.xpath('//body')[0]
        StructuringParser.fix_row_and_col_span(body)
        data = StructuringParser.handle_html_node(body)
        return data

    # @staticmethod
    # def handle_html_node(current_node):
    #     # 该节点所含文本
    #     string = current_node.text
    #     # 该节点包含子节点的所有的文本
    #     text = str(etree.tostring(current_node, encoding="utf8", method="text"), encoding="utf8")
    #     if len(list(current_node)) == 0 and string is not None and not string.isspace():
    #         # 叶子节点
    #         #LxmlHelper.get_text_of_node(current_node)
    #         return StringHelper.filter_illegal_char(text)
    #     elif text is None or len(text) == 0 or text.isspace():
    #         # 无效节点
    #         return None
    #     else:
    #         # 非叶子节点
    #         children = list(current_node)
    #         result = []
    #         if len(children) == 1:
    #             return StructuringParser.handle_html_node(children[0])
    #         elif len(children) == 0:
    #             return None
    #         for child in children:
    #             child_result = StructuringParser.handle_html_node(child)
    #             if child_result is not None:
    #                 result.append(child_result)
    #         if len(result) > 1:
    #             return result
    #         elif len(result) == 1:
    #             return result[0]
    #         return None

    @staticmethod
    def handle_html_node(current_node):
        string = current_node.text
        if len(list(current_node)) == 0 or (string is not None and not string.isspace()):
            # 叶子节点
            return StringHelper.filter_illegal_char(LxmlHelper.get_text_of_node(current_node))
        else:
            # 非叶子节点
            children = list(current_node)
            result = []
            if len(children) <= 1:
                return StructuringParser.handle_html_node(children[0])
            for child in children:
                result.append(StructuringParser.handle_html_node(child))
            return result

    # 递归处理Html节点的rowspan和colspan属性
    @staticmethod
    def fix_row_and_col_span(node):
        parent = node.getparent()
        index = parent.index(node)
        row_span = LxmlHelper.get_attribute_of_element(node, "rowspan")
        if row_span is not None:
            siblings = parent
            for i in range(0, int(row_span)-1):
                siblings = siblings.getnext()
                if siblings is None:
                    break
                copy_node = deepcopy(node)
                copy_node.set("rowspan", "0")
                siblings.insert(index, copy_node)
        # col_span = LxmlHelper.get_attribute_of_element(node, "colspan")
        # if col_span is not None:
        #     for i in range(0, int(col_span)-1):
        #         copy_node = deepcopy(node)
        #         copy_node.set("colspan", "0")
        #         parent.insert(index, copy_node)
        children = list(node)
        for child in children:
            StructuringParser.fix_row_and_col_span(child)

    @staticmethod
    def parse_docx(docx_file):
        docx = Document(docx_file)
        tables = docx.tables
        return StructuringParser.__handle_docx_tables(tables)

    @staticmethod
    def __handle_docx_tables(tables):
        result = []
        for i in range(0, len(tables)):
            table = tables[i]
            table_data = []
            for j in range(0, len(table.rows)):
                row = table.rows[j]
                row_data = []
                for k in range(0, len(row.cells)):
                    cell_tables = row.cells[k].tables
                    if len(cell_tables) != 0:
                        row_data.append(StructuringParser.__handle_docx_tables(cell_tables))
                    else:
                        row_data.append(row.cells[k].text)
                table_data.append(row_data)
            result.append(table_data)
        return result

    @staticmethod
    def parse_xls(xls_file):
        work_book = xlrd.open_workbook(xls_file)
        result = []
        for book in work_book.sheets():
            table_data = []
            for row_index in range(0, book.nrows):
                row_data = []
                for col_index in range(0, book.ncols):
                    row_data.append(StringHelper.filter_illegal_char(book.cell(row_index, col_index).value))
                table_data.append(row_data)
            result.append(table_data)
        return result


class DataMapper:
    class ColDataMapper:
        # 列映射器
        def __init__(self):
            super().__init__()
            self.attributes_map = OrderedDict()  # 属性映射{"源属性":"目的属性"}
            self.col_mapping_dic = self.attributes_map  # 列属性映射{"源属性":["目的属性"]}
            self.select_node = []  # 选择映射节点

        # 根据传入的col节点初始化配置
        # 当前的策略是将包含attri属性的col节点的attri属性的值作为源属性，col节点text内容作为目的属性
        def init_setting(self, col_node):
            attributes = col_node.attrib
            select_str = "select"
            if select_str in attributes.keys():
                self.select_node.append(col_node)
            else:
                for attribute in attributes.keys():
                    key = attributes[attribute]
                    if key in self.col_mapping_dic.keys():
                        self.col_mapping_dic[key].append(col_node.text)
                    else:
                        self.col_mapping_dic[key] = [col_node.text]

        def get_target_attributes(self):
            result = []
            for value in self.col_mapping_dic.values():
                result.extend(value)
            return ArrayHelper.remove_duplicate_item(result)

        def get_headers(self):
            result = []
            result.extend(self.get_target_attributes())
            for node in self.select_node:
                result.append(node.text)
            return ArrayHelper.remove_duplicate_item(result)

        def mapping(self, data):
            # 计算所有属性行的位置
            attributes_row_index_counter = {}
            for key in self.col_mapping_dic:
                coordinates = ArrayHelper.locate_value(data, key)
                for coordinate in coordinates:
                    x = tuple(coordinate[0:len(coordinate) - 1])
                    if x in attributes_row_index_counter.keys():
                        counter = attributes_row_index_counter[x]
                        attributes_row_index_counter[x] = counter + 1
                    else:
                        attributes_row_index_counter[x] = 1
            attributes_row_index = []
            number_of_target_attributes = len(self.get_target_attributes())
            for key in attributes_row_index_counter:
                if attributes_row_index_counter[key] >= min(number_of_target_attributes, 3):
                    attributes_row_index.append(key)
            print("所有可能的属性行如下:")
            for index in attributes_row_index:
                print(ArrayHelper.get_value_by_coordinate(data, index))

            # 映射
            last_data_row_index = -1  # 记录上一个数据区的最后一行下标
            for coordinate in attributes_row_index:
                attribute_row_index = coordinate[len(coordinate) - 1]
                start_row_index = attribute_row_index + 1
                attribute_row = ArrayHelper.get_value_by_coordinate(data, coordinate)
                table_coordinate = coordinate[0:len(coordinate) - 1]
                table = ArrayHelper.get_value_by_coordinate(data, table_coordinate)
                if table is None:
                    continue

                # 计算下标和属性的对应关系
                index_to_attribute = {}
                select_attributes = []
                for i in range(0, len(attribute_row)):
                    target_attribute = DataMapper.get_target_attribute(self.col_mapping_dic, attribute_row[i])
                    if target_attribute is not None:
                        index_to_attribute[i] = target_attribute
                        select_attributes.append(attribute_row[i])

                for i in range(start_row_index, len(table)):
                    row = table[i]
                    current_coordinate = list(coordinate[0:len(coordinate) - 1])
                    current_coordinate.append(i)
                    if len(row) != len(attribute_row) or tuple(current_coordinate) in attributes_row_index:
                        last_data_row_index = i - 1
                        break
                    result = OrderedDict()
                    for j in range(0, len(row)):
                        tmp = row[j]
                        if isinstance(tmp, list):
                            tmp = ''.join(tmp)
                        if j in index_to_attribute.keys():
                            for attribute in index_to_attribute[j]:
                                result[attribute] = tmp

                    # 处理select节点
                    for node in self.select_node:
                        select_str = LxmlHelper.get_attribute_of_element(node, "select")
                        regex_str = LxmlHelper.get_attribute_of_element(node, "regex")
                        attribute_str = node.text

                        # 以'@'开头表示选择的是表头属性或者标题
                        value = None
                        if select_str.startswith('@'):
                            header_result = re.search('(?<=@h)[0-9]*', select_str)
                            #  匹配表头
                            if header_result is not None:
                                header_index = int(header_result.group())
                                if regex_str is not None and len(regex_str) != 0:
                                    search_result = re.search(regex_str, select_attributes[header_index])
                                    if search_result is None:
                                        continue
                                    value = search_result.group()
                                else:
                                    value = select_attributes[header_index]
                            title_result = re.search('(?<=@t)[0-9]*', select_str)
                            #  匹配标题
                            if title_result is not None:
                                title_index = int(title_result.group())
                                target_title_row_index = attribute_row_index - title_index
                                if target_title_row_index >= 0 and target_title_row_index > last_data_row_index:
                                    title = None
                                    for x in table[target_title_row_index]:
                                        if x is not None and len(x) != 0:
                                            title = x
                                            break
                                    if title is not None:
                                        if regex_str is not None and len(regex_str) != 0:
                                            search_result = re.search(regex_str, title)
                                            if search_result is None:
                                                continue
                                            value = search_result.group()
                                        else:
                                            value = title
                        else:
                            # TODO 其他选择方式
                            pass

                        result[attribute_str] = value
                    yield result

    def __init__(self, setting_node):
        self.mapper = None
        self.__parse_setting_file(setting_node)

    # 获取表头
    def get_headers(self):
        return self.mapper.get_headers()

    # 解析映射规则
    def __parse_setting_file(self, mapping_node):
        tables = mapping_node.xpath('//table')
        for table in tables:
            children = list(table)
            for child in children:
                if child.tag == 'col':
                    if self.mapper is None:
                        self.mapper = DataMapper.ColDataMapper()
                    self.mapper.init_setting(child)

    # 处理映射
    def mapping(self, data):
        return self.mapper.mapping(data)

    @staticmethod
    def get_target_attribute(dir, attri):
        if attri == '':
            return None
        x = attri
        if isinstance(attri, list):
            x = ''.join(attri)
        for key in dir.keys():
            if StringHelper.match_string(x, key):
                return dir[key]
        return None
